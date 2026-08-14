"""Upload persistence plus native/visual document routing."""

from pathlib import Path
from uuid import uuid4

from fastapi import HTTPException, UploadFile
from pypdf import PdfReader
from docx import Document

from app.config import config
from app.modules.file_inspection import inspect_file
from app.modules.gemini_converter import GeminiHtmlConverter
from app.modules.html_processing import html_to_text
from app.modules.native_extraction import TEXT_EXTENSIONS, extract_native
from app.schemas.ingestion_models import IngestionRoute, PreparedDocument


def validate_file(filename: str) -> bool:
    """Return whether ``filename`` has an allowed extension."""
    if not filename:
        return False
    return Path(filename).suffix.lower() in {
        extension.lower() for extension in config.ALLOWED_FILE_TYPES
    }


def extract_text_from_pdf(file_path: str) -> str:
    try:
        reader = PdfReader(file_path)
        return "\n\n".join(
            text for page in reader.pages if (text := page.extract_text())
        ).strip()
    except Exception as exc:
        raise ValueError(f"Cannot read PDF file: {exc}") from exc


def extract_text_from_txt(file_path: str) -> str:
    path = Path(file_path)
    try:
        return path.read_text(encoding="utf-8-sig").strip()
    except UnicodeDecodeError:
        # A small fallback for common Windows text files.
        return path.read_text(encoding="cp1252").strip()
    except OSError as exc:
        raise ValueError(f"Cannot read TXT file: {exc}") from exc


def extract_text_from_docx(file_path: str) -> str:
    try:
        document = Document(file_path)
        return "\n\n".join(
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ).strip()
    except Exception as exc:
        raise ValueError(f"Cannot read DOCX file: {exc}") from exc


async def save_uploaded_file(upload_file: UploadFile) -> str:
    """Stream an upload to disk while enforcing the configured size limit."""
    original_name = Path(upload_file.filename or "").name
    if not validate_file(original_name):
        allowed = ", ".join(config.ALLOWED_FILE_TYPES)
        raise HTTPException(400, f"Unsupported file type. Allowed: {allowed}")

    upload_dir = Path(config.UPLOAD_DIR)
    upload_dir.mkdir(parents=True, exist_ok=True)
    destination = upload_dir / f"{uuid4().hex}_{original_name}"
    total_size = 0

    try:
        await upload_file.seek(0)
        with destination.open("wb") as output:
            while content := await upload_file.read(1024 * 1024):
                total_size += len(content)
                if total_size > config.MAX_FILE_SIZE:
                    raise HTTPException(
                        413,
                        f"File exceeds {config.MAX_FILE_SIZE // (1024 * 1024)} MB limit",
                    )
                output.write(content)
    except Exception:
        destination.unlink(missing_ok=True)
        raise
    finally:
        await upload_file.close()

    if total_size == 0:
        destination.unlink(missing_ok=True)
        raise HTTPException(400, "Uploaded file is empty")
    return str(destination)


async def process_uploaded_file(upload_file: UploadFile) -> PreparedDocument:
    """Save, inspect, classify, and normalize an uploaded document."""
    original_name = Path(upload_file.filename or "").name
    file_path = await save_uploaded_file(upload_file)
    try:
        inspection = inspect_file(file_path, original_name)
        extraction = extract_native(inspection)
        if extraction.acceptable:
            return PreparedDocument(
                inspection=inspection,
                route=IngestionRoute.DIRECT_TEXT,
                text=extraction.text,
                extractor=extraction.extractor,
                warnings=extraction.warnings,
            )
        if inspection.extension in TEXT_EXTENSIONS:
            raise ValueError("No readable text could be extracted from the file")

        route = (
            IngestionRoute.HYBRID_HTML
            if extraction.has_partial_pages
            else IngestionRoute.GEMINI_HTML
        )
        conversion = await GeminiHtmlConverter().convert(inspection)
        return PreparedDocument(
            inspection=inspection,
            route=route,
            text=html_to_text(conversion.html),
            html=conversion.html,
            extractor="gemini",
            normalized_html_path=conversion.artifact_path,
            converter_model=conversion.model,
            prompt_version=conversion.prompt_version,
            cached_conversion=conversion.cached,
            warnings=extraction.warnings,
        )
    except Exception as exc:
        Path(file_path).unlink(missing_ok=True)
        if isinstance(exc, HTTPException):
            raise
        raise HTTPException(400, str(exc)) from exc
